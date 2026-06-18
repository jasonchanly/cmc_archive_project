from pypdf import PdfReader
import docx
import json
from pathlib import Path
import pandas as pd
import re
import calendar
import convertapi
from doc2docx import convert
import shutil
import argparse
from datetime import datetime

MONTH2NUM = {v: i for i, v in enumerate(calendar.month_name) if v} | {v: i for i, v in enumerate(calendar.month_abbr) if v}

def pdf2txt(pdf_fpath: str | Path):
    reader = PdfReader(pdf_fpath)
    pdf_text = ""
    for page in reader.pages:
        pdf_text += page.extract_text()
    return pdf_text

def docx2txt(doc_fpath: Path):
    if doc_fpath.suffix == ".doc":
        convert(fpath, "_temp.docx", keep_active=True)
    # workaround - first convert doc to docx, and then apply docx2txt converter
        doc = docx.Document("_temp.docx")
    else:
        doc = docx.Document(fpath)
    docx_text = "\n".join([para.text for para in doc.paragraphs])
    return docx_text


def autoconvert_programme(fpath: str, txt_dir: Path, ocr_dir: Path):
    """
    Reads a fpath - determines its extension, attempts to extract its text into a txt file:
    if fail (only expected for PDFs), preserves the file as a PDF.
    Then save it to the correct dir (dir for txt files vs dir for un-OCRed pdfs, respectively)
    :param fpath:
    :param txt_dir:
    :param ocr_dir:
    :return:
    """
    fpath = Path(fpath)
    txt_file_name = fpath.name.replace(fpath.suffix, ".txt")
    if (txt_dir/txt_file_name).exists():
        print("File already parsed")
        return
    elif (ocr_dir/fpath.name).exists():
        print("File already parsed")
        return

    if fpath.suffix == ".pdf":
        extracted_text = pdf2txt(fpath)
        if not extracted_text.strip():
            # if text extraction fail, copy file as is
            shutil.copy(fpath, ocr_dir/fpath.name)
            return

    elif fpath.suffix in (".doc", ".docx"):
        extracted_text = docx2txt(fpath)

    else:
        raise ValueError(f"Unsupported file extension: {fpath.suffix}")

    # convoluted, as fpath.stem

    with open(txt_dir/txt_file_name, "w") as txt_file:
        txt_file.write(extracted_text)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--filelist", type=str, help="Txt filepath containing the list of filepaths to process")
    parser.add_argument("--batch_name", type=str, required=False)
    args = parser.parse_args()

    problematic_filelist = args.filelist.replace(".txt", "_problematic_files.txt")

    batch_name = "batch_" + datetime.now().strftime("%Y_%m_%d") if not args.batch_name else args.batch_name

    ocr_dir = Path("data")/(batch_name + "_ocr_needed")
    txt_dir = Path("data")/(batch_name + "_no_ocr_needed")

    ocr_dir.mkdir(exist_ok=True)
    txt_dir.mkdir(exist_ok=True)

    with open(args.filelist, "r") as f:
        fpaths_list = [x.strip() for x in f.read().splitlines() if x.strip()]

    problematic_files = []

    for fpath in fpaths_list:
        print(f"Processing {fpath}")
        try:
            autoconvert_programme(fpath, txt_dir, ocr_dir)
        except:
            print(f"Problematic file to inspect: {fpath} - excluded for now")
            problematic_files.append(Path(fpath).name)

    print("Processing complete, verifying files")

    for fpath in txt_dir.glob("*.txt"):
        if fpath.stat().st_size < 400:
            problematic_files.append(fpath.name)

    with open(problematic_filelist, "w") as f:
        for fpath in problematic_files:
            f.write(fpath + "\n")
